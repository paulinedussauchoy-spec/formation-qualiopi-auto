"""
Tests de génération des certificats de réalisation — cas GENERFEU.

Vérifie :
  - 39 certificats générés (1 par stagiaire, TNS inclus)
  - Taille raisonnable (> 10 Ko)
  - Noms de fichiers cohérents (CERTIF_GENERFEU_...)
  - Pas de tags Jinja2 résiduels
  - Contenu du certificat DUPONT (Admins) : nom, prenom, client, module
  - Certificat BERTRAND (TNS Admins) : présent malgré exclusion de la convention
"""

import sys
from pathlib import Path

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

from core.excel_parser import parse_formations_excel
from core.module_mapper import ModuleMapper
from core.group_builder import build_convention_groups
from core.document_gen import CertificatRenderer, ClientInfo

EXCEL_PATH = ROOT / "base de travail" / "GENERFEU - Formations et inscrits.xlsx"
OUTPUT_DIR  = ROOT / "output" / "test_certificats"

GENERFEU = ClientInfo(
    nom="GENERFEU",
    adresse="PARC D'ACTIVITES DES ECLAPONS\n3 CHEMIN DES ECLAPONS\n69390 VOURLES",
    representant="Gilles BERTRAND",
    fonction="Dirigeant",
    frais_mission_ht=0.0,
)

DATES    = "Du 01/04/2026 au 31/10/2026"
DATE_DOC = "01/03/2026"


def _load():
    data = parse_formations_excel(EXCEL_PATH)
    mapper = ModuleMapper()
    groups, _ = build_convention_groups(data, mapper)
    renderer = CertificatRenderer()
    return groups, renderer


def _generate_all(groups, renderer):
    return renderer.generate_all(
        groups=groups,
        client=GENERFEU,
        dates_previsionnelles=DATES,
        date_document=DATE_DOC,
        ref_dossier="2026-GENERFEU",
        output_dir=OUTPUT_DIR,
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_genere_39_certificats():
    """39 stagiaires au total (TNS inclus) → 39 certificats."""
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    assert len(paths) == 39, (
        f"Attendu 39 certificats, obtenu {len(paths)}\n"
        + "\n".join(f"  {p.name}" for p in paths)
    )
    print(f"  {len(paths)} certificats générés — OK")


def test_certificats_non_vides():
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    for p in paths:
        size = p.stat().st_size
        assert size > 10_000, f"{p.name} trop petit : {size} octets"
    print(f"  Tous les certificats > 10 Ko — OK")


def test_noms_fichiers_certif():
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    for p in paths:
        assert p.name.startswith("CERTIF_GENERFEU_"), (
            f"Nom inattendu : {p.name}"
        )
        assert p.suffix == ".docx"
    print("  Noms de fichiers CERTIF_GENERFEU_... — OK")


def test_pas_de_tags_residuels_certif():
    """Aucun tag Jinja2 {{ }} ou {%tr %} ne doit subsister."""
    import re
    from docx import Document

    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    pattern = re.compile(r"\{\{|\}\}|\{%|%\}")
    for p in paths:
        doc = Document(p)
        texts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        full = " ".join(texts)
        matches = pattern.findall(full)
        assert not matches, f"Tags résiduels dans {p.name} : {set(matches)}"
    print("  Aucun tag Jinja2 résiduel — OK")


def test_contenu_certif_dupont():
    """Vérifie le certificat de DUPONT (Administrateurs)."""
    from docx import Document

    groups, renderer = _load()
    _generate_all(groups, renderer)

    certif = OUTPUT_DIR / "CERTIF_GENERFEU_DUPONT_Christelle_Administrateurs.docx"
    assert certif.exists(), f"Certificat DUPONT introuvable : {certif}"

    doc = Document(certif)
    full_text = " ".join(
        para.text for para in doc.paragraphs
    ) + " " + " ".join(
        cell.text for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )

    assert "DUPONT" in full_text,              "NOM DUPONT manquant"
    assert "Christelle" in full_text,          "Prénom Christelle manquant"
    assert "GENERFEU" in full_text,            "Nom client manquant"
    assert "ECLAPONS" in full_text,            "Adresse client manquante"
    assert "1-01" in full_text,               "Module 1-01 manquant"
    assert "1-03" in full_text,               "Module 1-03 manquant"
    assert "14 heures" in full_text,          "Durée 14 heures manquante"
    assert "2026-GENERFEU" in full_text,       "Référence dossier manquante"
    assert "01/03/2026" in full_text,          "Date document manquante"
    assert "À distance" in full_text,          "Modalité manquante"
    assert "Du 01/04/2026" in full_text,       "Dates prévisionnelles manquantes"
    print("  Certificat DUPONT Christelle : tous les champs vérifiés — OK")


def test_certif_bertrand_tns_present():
    """BERTRAND (TNS, exclu de la convention) doit avoir un certificat."""
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)

    bertrand_paths = [p for p in paths if "BERTRAND" in p.name]
    assert len(bertrand_paths) == 1, (
        f"Attendu 1 certif BERTRAND (TNS), obtenu {len(bertrand_paths)}"
    )
    print(f"  BERTRAND TNS : certificat présent → {bertrand_paths[0].name}")


def test_dates_realisees_injectees():
    """Les dates réelles passées via dates_realisees_per_group s'affichent sur le certificat."""
    from docx import Document

    groups, renderer = _load()

    # Simuler les dates réelles pour les Admins (groupe 0)
    dates_admins = "01/04/2026, 08/04/2026, 15/04/2026, 22/04/2026"
    dates_per_group = [dates_admins] + [None] * (len(groups) - 1)

    paths = renderer.generate_all(
        groups=groups,
        client=GENERFEU,
        dates_previsionnelles=DATES,
        date_document=DATE_DOC,
        ref_dossier="2026-GENERFEU",
        output_dir=OUTPUT_DIR,
        dates_realisees_per_group=dates_per_group,
    )

    certif = OUTPUT_DIR / "CERTIF_GENERFEU_DUPONT_Christelle_Administrateurs.docx"
    assert certif.exists()
    doc = Document(certif)
    full_text = " ".join(
        cell.text for table in doc.tables
        for row in table.rows for cell in row.cells
    )
    assert "01/04/2026, 08/04/2026" in full_text, (
        f"Dates réalisées non trouvées dans le certificat. Contenu : {full_text[:300]}"
    )
    assert "Dates de r\u00e9alisation" in full_text, "Label 'Dates de réalisation' manquant"
    print(f"  Dates réalisées injectées : {dates_admins} — OK")


def test_certif_techniciens_17():
    """17 certificats pour le groupe Techniciens."""
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)

    # Le groupe Techniciens a le label contenant "1-10"
    tech_paths = [p for p in paths if "1-10" in p.name and "1-02" not in p.name and "2-06" not in p.name]
    assert len(tech_paths) == 17, (
        f"Attendu 17 certificats Techniciens, obtenu {len(tech_paths)}\n"
        + "\n".join(f"  {p.name}" for p in tech_paths)
    )
    print(f"  Techniciens : {len(tech_paths)} certificats — OK")


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

def run_all():
    tests = [
        test_genere_39_certificats,
        test_certificats_non_vides,
        test_noms_fichiers_certif,
        test_pas_de_tags_residuels_certif,
        test_contenu_certif_dupont,
        test_dates_realisees_injectees,
        test_certif_bertrand_tns_present,
        test_certif_techniciens_17,
    ]

    passed = failed = 0
    for test in tests:
        name = test.__name__
        try:
            test()
            print(f"PASS  {name}")
            passed += 1
        except AssertionError as e:
            print(f"FAIL  {name}: {e}")
            failed += 1
        except Exception as e:
            import traceback
            print(f"ERROR {name}: {type(e).__name__}: {e}")
            traceback.print_exc()
            failed += 1

    print(f"\n{passed}/{passed + failed} tests passés")
    return failed == 0


if __name__ == "__main__":
    success = run_all()
    sys.exit(0 if success else 1)
