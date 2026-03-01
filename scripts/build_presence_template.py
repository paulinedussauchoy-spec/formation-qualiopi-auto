"""
build_presence_template.py — Génère templates/feuille_presence.docx
depuis "base de travail/[CLIENT] FEUILLE DE PRESENCE.docx".

À lancer UNE FOIS (ou après modification du document source).

Logique métier :
  1 feuille par demi-journée (= par colonne Excel / par module_column)
  Chaque feuille couvre 1 seul module (pas de boucle pour Table 1)
  Stats fixes : 0,5 journée / 3,5 heures
  Stagiaires : tous (TNS inclus, car présents lors des sessions)

Tables modifiées :
  Table 0 (1 col)  — Nom et adresse entreprise  → remplacement simple
  Table 1 (1 col)  — Nom de la formation (1 seul module) → remplacement simple
  Table 2 (2 cols) — stats session               → remplacement simple
  Table 3 (4 cols) — liste stagiaires + signatures → boucle {%tr for s in stagiaires %}

Paragraphes modifiés :
  "Numéro de référence"  → {{ ref_dossier }}
  "Date :"               → {{ date_session }}
"""

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent.parent
SRC  = ROOT / "base de travail" / "[CLIENT] FEUILLE DE PRESENCE.docx"
DEST = ROOT / "templates" / "feuille_presence.docx"


# ---------------------------------------------------------------------------
# Helpers XML (identiques aux autres scripts build_*)
# ---------------------------------------------------------------------------

def _get_tc_list(tr_element):
    return tr_element.findall(qn("w:tc"))


def _clear_tc(tc):
    paras = tc.findall(qn("w:p"))
    for p in paras[1:]:
        tc.remove(p)
    para = paras[0]
    for r in para.findall(qn("w:r")):
        para.remove(r)
    return para


def _set_tc_text(tc, text: str):
    para = _clear_tc(tc)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    para.append(r)


def _set_cell(cell, text: str):
    _set_tc_text(cell._tc, text)


def _delete_rows_from(table, from_idx: int):
    tbl = table._tbl
    rows = table.rows
    for i in range(len(rows) - 1, from_idx - 1, -1):
        tbl.remove(rows[i]._tr)


def _clone_row_with_texts(table, source_row_idx: int, texts: list):
    source_tr = table.rows[source_row_idx]._tr
    new_tr = deepcopy(source_tr)
    tcs = _get_tc_list(new_tr)
    for i, tc in enumerate(tcs):
        text = texts[i] if i < len(texts) else ""
        _set_tc_text(tc, text)
    table._tbl.append(new_tr)


# ---------------------------------------------------------------------------
# Construction du template
# ---------------------------------------------------------------------------

def build_presence_template():
    if not SRC.exists():
        print(f"ERREUR : document source introuvable : {SRC}", file=sys.stderr)
        sys.exit(1)

    DEST.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(SRC, DEST)
    doc = Document(DEST)

    # ------------------------------------------------------------------
    # TABLE 0 — Nom et adresse de l'entreprise (2 lignes × 1 colonne)
    # ------------------------------------------------------------------
    t0 = doc.tables[0]
    # Ligne 0 : "Nom et Adresse de l'entreprise" (conserver)
    _set_cell(t0.rows[1].cells[0], "{{ client_nom }}\n{{ client_adresse }}")

    # ------------------------------------------------------------------
    # TABLE 1 — Nom de la formation (2 lignes × 1 colonne)
    #
    # 1 seul module par feuille (pas de boucle) :
    #   Ligne 0 : "Nom de la formation" (conserver)
    #   Ligne 1 : {{ module_code }} - {{ module_intitule }}
    # ------------------------------------------------------------------
    t1 = doc.tables[1]
    # Ligne 0 : en-tête (conserver)
    _set_cell(t1.rows[1].cells[0], "{{ module_code }} - {{ module_intitule }}")

    # ------------------------------------------------------------------
    # TABLE 2 — Statistiques session (5 lignes × 2 colonnes)
    #
    # Stats fixes pour 1 demi-journée :
    #   nb_journees_str = "0,5"   nb_heures_str = "3,5 heures"
    # ------------------------------------------------------------------
    t2 = doc.tables[2]
    _set_cell(t2.rows[0].cells[1], "{{ nb_journees_str }}")
    _set_cell(t2.rows[1].cells[1], "{{ nb_heures_str }}")
    _set_cell(t2.rows[2].cells[1], "{{ modalite }}")
    # Ligne 3 "Formation de type collectif" — laisser tel quel
    _set_cell(t2.rows[4].cells[1], "{{ dates_session_detail }}")

    # ------------------------------------------------------------------
    # TABLE 3 — Stagiaires + signatures (9 lignes × 4 colonnes → boucle)
    #
    # Structure correcte pour docxtpl :
    #   Ligne for     : {%tr for s in stagiaires %}  ← ligne tag seule (vide)
    #   Ligne content : {{ s.nom }} | {{ s.prenom }} ← corps répété N fois
    #   Ligne endfor  : {%tr endfor %}               ← ligne tag seule
    # ------------------------------------------------------------------
    t3 = doc.tables[3]
    # Ligne 0 : en-têtes NOM | Prénom | Signature stagiaires | Signature formateur (conserver)
    # Ligne 1 → tag de boucle seul (4 cellules vides)
    for ci in range(len(t3.rows[1].cells)):
        text = "{%tr for s in stagiaires %}" if ci == 0 else ""
        _set_cell(t3.rows[1].cells[ci], text)
    # Supprimer lignes 2-8 (lignes vides du template source)
    _delete_rows_from(t3, from_idx=2)
    # Ajouter ligne content : NOM | Prénom | (vide signature) | (vide formateur)
    _clone_row_with_texts(t3, source_row_idx=1,
                          texts=["{{ s.nom }}", "{{ s.prenom }}", "", ""])
    # Ajouter ligne endfor
    _clone_row_with_texts(t3, source_row_idx=1,
                          texts=["{%tr endfor %}", "", "", ""])

    # ------------------------------------------------------------------
    # PARAGRAPHES
    # ------------------------------------------------------------------

    # "Numéro de référence (si disponible) :" → {{ ref_dossier }}
    for p in doc.paragraphs:
        if p.text.strip().startswith("Num\u00e9ro de r\u00e9f\u00e9rence"):
            if p.runs:
                p.runs[0].text = "Num\u00e9ro de r\u00e9f\u00e9rence (si disponible)\xa0: {{ ref_dossier }}"
            break

    # "Date :" → "Date : {{ date_session }}"
    for p in doc.paragraphs:
        if p.text.strip() == "Date\xa0:":
            if p.runs:
                p.runs[0].text = "Date\xa0: {{ date_session }}"
            break

    # ------------------------------------------------------------------
    # Sauvegarde
    # ------------------------------------------------------------------
    doc.save(DEST)
    print(f"Template créé : {DEST}")
    _verify_template(doc)


def _verify_template(doc):
    import re
    found = set()
    pattern = re.compile(r"\{\{[^}]+\}\}|\{%[^%]+%\}")

    for p in doc.paragraphs:
        for m in pattern.finditer(p.text):
            found.add(m.group())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for m in pattern.finditer(cell.text):
                    found.add(m.group())

    print(f"\nTags Jinja2 trouvés ({len(found)}) :")
    for tag in sorted(found):
        print(f"  {tag}")


if __name__ == "__main__":
    build_presence_template()
