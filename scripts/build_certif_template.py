"""
build_certif_template.py — Génère templates/certificat.docx depuis [CLIENT] - CERTIF REALISATION.docx

À lancer UNE FOIS (ou après modification du document source).
Injecte les tags Jinja2/docxtpl en conservant la mise en forme du document original.

Tables modifiées :
  Table 0 (2 cols) — NOM | Prénom du stagiaire    → remplacement simple
  Table 1 (1 col)  — Nom et adresse entreprise    → remplacement simple
  Table 2 (1 col)  — liste modules                → boucle {%tr for m in modules %}
  Table 3 (2 cols) — stats (journées/heures/lieu) → remplacement simple

Paragraphes modifiés :
  "Références du dossier :"  → ajout {{ ref_dossier }}
  "Le :"                     → ajout {{ date_document }}
"""

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent.parent
SRC  = ROOT / "base de travail" / "[CLIENT] - CERTIF REALISATION.docx"
DEST = ROOT / "templates" / "certificat.docx"


# ---------------------------------------------------------------------------
# Helpers XML (identiques à build_templates.py)
# ---------------------------------------------------------------------------

def _get_tc_list(tr_element):
    return tr_element.findall(qn("w:tc"))


def _clear_tc(tc):
    paras = tc.findall(qn("w:p"))
    for p in paras[1:]:
        tc.remove(p)
    para = paras[0]
    for r in para.findall(qn("w:r")):
        para.remove(r)
    return para


def _set_tc_text(tc, text: str):
    para = _clear_tc(tc)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    para.append(r)


def _set_cell(cell, text: str):
    _set_tc_text(cell._tc, text)


def _delete_rows_from(table, from_idx: int):
    tbl = table._tbl
    rows = table.rows
    for i in range(len(rows) - 1, from_idx - 1, -1):
        tbl.remove(rows[i]._tr)


def _clone_row_with_texts(table, source_row_idx: int, texts: list):
    source_tr = table.rows[source_row_idx]._tr
    new_tr = deepcopy(source_tr)
    tcs = _get_tc_list(new_tr)
    for i, tc in enumerate(tcs):
        text = texts[i] if i < len(texts) else ""
        _set_tc_text(tc, text)
    table._tbl.append(new_tr)


# ---------------------------------------------------------------------------
# Construction du template
# ---------------------------------------------------------------------------

def build_certif_template():
    if not SRC.exists():
        print(f"ERREUR : document source introuvable : {SRC}", file=sys.stderr)
        sys.exit(1)

    DEST.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(SRC, DEST)
    doc = Document(DEST)

    # ------------------------------------------------------------------
    # TABLE 0 — NOM | Prénom du stagiaire (2 lignes × 2 colonnes)
    # ------------------------------------------------------------------
    t0 = doc.tables[0]
    # Ligne 0 : en-têtes "NOM" | "Prénom" (conserver)
    _set_cell(t0.rows[1].cells[0], "{{ stagiaire_nom }}")
    _set_cell(t0.rows[1].cells[1], "{{ stagiaire_prenom }}")

    # ------------------------------------------------------------------
    # TABLE 1 — Nom et adresse de l'entreprise (2 lignes × 1 colonne)
    # ------------------------------------------------------------------
    t1 = doc.tables[1]
    # Ligne 0 : "Nom et Adresse de l'entreprise" (conserver)
    _set_cell(t1.rows[1].cells[0], "{{ client_nom }}\n{{ client_adresse }}")

    # ------------------------------------------------------------------
    # TABLE 2 — Modules suivis (5 lignes × 1 colonne → boucle)
    #
    # Structure correcte pour docxtpl :
    #   Ligne for     : {%tr for m in modules %}   ← remplacée par {% for m in modules %}
    #   Ligne content : {{ m.code }} - ...          ← corps répété N fois
    #   Ligne endfor  : {%tr endfor %}              ← remplacée par {% endfor %}
    # ------------------------------------------------------------------
    t2 = doc.tables[2]
    # Ligne 0 : en-tête "Nom de la formation" (conserver)
    # Ligne 1 → tag de boucle seul
    _set_cell(t2.rows[1].cells[0], "{%tr for m in modules %}")
    # Supprimer lignes 2-4 (vides dans le template source)
    _delete_rows_from(t2, from_idx=2)
    # Ajouter ligne content
    _clone_row_with_texts(t2, source_row_idx=1, texts=["{{ m.code }} - {{ m.intitule }}"])
    # Ajouter ligne endfor
    _clone_row_with_texts(t2, source_row_idx=1, texts=["{%tr endfor %}"])

    # ------------------------------------------------------------------
    # TABLE 3 — Statistiques (5 lignes × 2 colonnes)
    # ------------------------------------------------------------------
    t3 = doc.tables[3]
    _set_cell(t3.rows[0].cells[1], "{{ nb_journees_str }}")
    _set_cell(t3.rows[1].cells[1], "{{ nb_heures_str }}")
    _set_cell(t3.rows[2].cells[1], "{{ modalite }}")
    # Ligne 3 "Formation de type collectif" — laisser tel quel
    # Ligne 4 : label → "Dates de réalisation", valeur → dates effectives
    _set_cell(t3.rows[4].cells[0], "Dates de r\u00e9alisation")
    _set_cell(t3.rows[4].cells[1], "{{ dates_realisees }}")
    # Ligne 5 (nouvelle) : horaires des sessions
    _clone_row_with_texts(t3, source_row_idx=4, texts=["Horaires des sessions", "{{ horaires_sessions }}"])

    # ------------------------------------------------------------------
    # PARAGRAPHES
    # ------------------------------------------------------------------

    # "Références du dossier :" → ajout {{ ref_dossier }}
    for p in doc.paragraphs:
        if p.text.strip().startswith("R\u00e9f\u00e9rences du dossier"):
            if p.runs:
                p.runs[0].text = "R\u00e9f\u00e9rences du dossier\xa0: {{ ref_dossier }}"
            break

    # "Le :" → ajout {{ date_document }} (run[5] du paragraphe "Fait à / Le :")
    for p in doc.paragraphs:
        if "Fait \u00e0" in p.text and "Le :" in p.text:
            for r in p.runs:
                if r.text.strip() == "Le :":
                    r.text = "\nLe : {{ date_document }}"
                    break
            break

    # ------------------------------------------------------------------
    # Sauvegarde
    # ------------------------------------------------------------------
    doc.save(DEST)
    print(f"Template créé : {DEST}")
    _verify_template(doc)


def _verify_template(doc):
    import re
    found = set()
    pattern = re.compile(r"\{\{[^}]+\}\}|\{%[^%]+%\}")

    for p in doc.paragraphs:
        for m in pattern.finditer(p.text):
            found.add(m.group())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for m in pattern.finditer(cell.text):
                    found.add(m.group())

    print(f"\nTags Jinja2 trouvés ({len(found)}) :")
    for tag in sorted(found):
        print(f"  {tag}")


if __name__ == "__main__":
    build_certif_template()
