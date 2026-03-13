"""
Tests de génération des feuilles de présence — cas GENERFEU.

Logique : 1 feuille par demi-journée = 1 feuille par module_column.

Attendu pour GENERFEU :
  Administrateurs  (cols 5-8)   → 4 colonnes → 4 feuilles
  Utilisateurs Gr01 (cols 9-12) → 4 colonnes → 4 feuilles
  Utilisateurs Gr02 (cols 13-14)→ 2 colonnes → 2 feuilles
  SAV Gr03 (cols 15-17)         → 3 colonnes → 3 feuilles
  Techniciens (cols 18-19)      → 2 colonnes → 2 feuilles
  TOTAL                         = 15 feuilles

Vérifie :
  - 15 fichiers générés
  - Taille raisonnable (> 10 Ko)
  - Noms de fichiers cohérents (PRESENCE_GENERFEU_G{col}_...)
  - Pas de tags Jinja2 résiduels
  - Contenu : module correct, stagiaires présents, stats demi-journée
  - Feuille Admins DJ01 → module 1-01, 6 stagiaires (incl. BERTRAND TNS)
  - Feuille Techniciens → 17 stagiaires, module 1-10
"""

import sys
from pathlib import Path

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

from core.excel_parser import parse_formations_excel
from core.module_mapper import ModuleMapper
from core.group_builder import build_convention_groups
from core.document_gen import FeuillePresenceRenderer, ClientInfo

EXCEL_PATH = ROOT / "base de travail" / "GENERFEU - Formations et inscrits.xlsx"
OUTPUT_DIR  = ROOT / "output" / "test_presence"

GENERFEU = ClientInfo(
    nom="GENERFEU",
    adresse="PARC D'ACTIVITES DES ECLAPONS\n3 CHEMIN DES ECLAPONS\n69390 VOURLES",
    representant="Gilles BERTRAND",
    fonction="Dirigeant",
    frais_mission_ht=0.0,
)

DATES    = "Du 01/04/2026 au 31/10/2026"
DATE_DOC = "01/03/2026"


def _load():
    data = parse_formations_excel(EXCEL_PATH)
    mapper = ModuleMapper()
    groups, _ = build_convention_groups(data, mapper)
    renderer = FeuillePresenceRenderer()
    return groups, renderer


def _generate_all(groups, renderer):
    return renderer.generate_all(
        groups=groups,
        client=GENERFEU,
        dates_previsionnelles=DATES,
        date_document=DATE_DOC,
        ref_dossier="2026-GENERFEU",
        output_dir=OUTPUT_DIR,
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_genere_15_feuilles():
    """4+4+2+3+2 = 15 feuilles pour GENERFEU."""
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    assert len(paths) == 15, (
        f"Attendu 15 feuilles, obtenu {len(paths)}\n"
        + "\n".join(f"  {p.name}" for p in paths)
    )
    print(f"  {len(paths)} feuilles de présence générées — OK")


def test_feuilles_non_vides():
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    for p in paths:
        size = p.stat().st_size
        assert size > 10_000, f"{p.name} trop petit : {size} octets"
    print(f"  Toutes les feuilles > 10 Ko — OK")


def test_noms_fichiers_presence():
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    for p in paths:
        assert p.name.startswith("PRESENCE_GENERFEU_G"), (
            f"Nom inattendu : {p.name}"
        )
        assert p.suffix == ".docx"
        # Doit contenir _DJ suivi de 2 chiffres
        assert "_DJ" in p.name, f"Numéro DJ manquant dans : {p.name}"
    print("  Noms PRESENCE_GENERFEU_G{col}_..._DJ{nn}.docx — OK")
    for p in paths:
        print(f"    {p.name}")


def test_pas_de_tags_residuels():
    import re
    from docx import Document

    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    pattern = re.compile(r"\{\{|\}\}|\{%|%\}")
    for p in paths:
        doc = Document(p)
        texts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        full = " ".join(texts)
        matches = pattern.findall(full)
        assert not matches, f"Tags résiduels dans {p.name} : {set(matches)}"
    print("  Aucun tag Jinja2 résiduel — OK")


def test_feuille_admins_dj01():
    """Feuille Admins DJ1 : module 1-01, 6 stagiaires (BERTRAND TNS inclus), stats 0,5 journée."""
    from docx import Document

    groups, renderer = _load()
    _generate_all(groups, renderer)

    feuille = OUTPUT_DIR / "PRESENCE_GENERFEU_G05_Administrateurs_DJ01.docx"
    assert feuille.exists(), f"Feuille introuvable : {feuille}"

    doc = Document(feuille)
    full_text = " ".join(
        para.text for para in doc.paragraphs
    ) + " " + " ".join(
        cell.text for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )

    # Module correct pour DJ01 (col 5 = Admins + Socle commun → 1-01)
    assert "1-01" in full_text,           "Module 1-01 manquant"
    # Client
    assert "GENERFEU" in full_text,       "Nom client manquant"
    assert "ECLAPONS" in full_text,       "Adresse client manquante"
    # Stats demi-journée
    assert "0,5" in full_text,            "0,5 journée manquant"
    assert "3,5" in full_text,            "3,5 heures manquant"
    # Modalité Admins = VISIO → "À distance"
    assert "\u00c0 distance" in full_text, "Modalité À distance manquante"
    # Référence dossier
    assert "2026-GENERFEU" in full_text,  "Référence dossier manquante"
    # Stagiaires
    assert "DUPONT" in full_text,         "DUPONT manquant"
    assert "BERTRAND" in full_text,       "BERTRAND (TNS) manquant dans la feuille de présence"
    print("  Feuille Admins DJ01 (1-01, 0,5j, À distance, BERTRAND TNS) — OK")


def test_feuille_admins_dj02_module_1_03():
    """Feuille Admins DJ2 : col 6 = CRM → module 1-03."""
    from docx import Document

    groups, renderer = _load()
    _generate_all(groups, renderer)

    feuille = OUTPUT_DIR / "PRESENCE_GENERFEU_G05_Administrateurs_DJ02.docx"
    assert feuille.exists(), f"Feuille introuvable : {feuille}"

    doc = Document(feuille)
    full_text = " ".join(
        cell.text for table in doc.tables
        for row in table.rows for cell in row.cells
    )
    assert "1-03" in full_text, "Module 1-03 manquant (DJ02 Admins = CRM)"
    print("  Feuille Admins DJ02 (1-03 CRM) — OK")


def test_feuilles_admins_count():
    """4 feuilles pour le groupe Administrateurs."""
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    admin_paths = [p for p in paths if "G05_Administrateurs" in p.name]
    assert len(admin_paths) == 4, (
        f"Attendu 4 feuilles Admins, obtenu {len(admin_paths)}"
    )
    print(f"  Administrateurs : {len(admin_paths)} feuilles (DJ01-DJ04) — OK")


def test_feuilles_gr01_count_4():
    """4 feuilles pour Gr01 (4 colonnes, 2 modules répétés)."""
    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    gr01_paths = [p for p in paths if "G09_" in p.name]
    assert len(gr01_paths) == 4, (
        f"Attendu 4 feuilles Gr01, obtenu {len(gr01_paths)}"
    )
    print(f"  Gr01 (cols 9-12) : {len(gr01_paths)} feuilles — OK")


def test_feuilles_techniciens_17_stagiaires():
    """Feuilles Techniciens : 17 stagiaires par feuille, module 1-10."""
    from docx import Document

    groups, renderer = _load()
    paths = _generate_all(groups, renderer)
    tech_paths = [p for p in paths if "G18_" in p.name]
    assert len(tech_paths) == 2, (
        f"Attendu 2 feuilles Techniciens, obtenu {len(tech_paths)}"
    )

    for p in tech_paths:
        doc = Document(p)
        # Table 3 = stagiaires (row 0 header, dernière row = endfor vide)
        t3 = doc.tables[3]
        stagiaire_rows = [
            row for row in t3.rows[1:]
            if row.cells[0].text.strip()
        ]
        assert len(stagiaire_rows) == 17, (
            f"{p.name} : attendu 17 stagiaires, obtenu {len(stagiaire_rows)}"
        )
        full_text = " ".join(
            cell.text for row in t3.rows for cell in row.cells
        )
        assert "1-10" in " ".join(
            cell.text for table in doc.tables
            for row in table.rows for cell in row.cells
        ), f"Module 1-10 manquant dans {p.name}"

    print(f"  Techniciens : 2 feuilles × 17 stagiaires, module 1-10 — OK")


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

def run_all():
    tests = [
        test_genere_15_feuilles,
        test_feuilles_non_vides,
        test_noms_fichiers_presence,
        test_pas_de_tags_residuels,
        test_feuille_admins_dj01,
        test_feuille_admins_dj02_module_1_03,
        test_feuilles_admins_count,
        test_feuilles_gr01_count_4,
        test_feuilles_techniciens_17_stagiaires,
    ]

    passed = failed = 0
    for test in tests:
        name = test.__name__
        try:
            test()
            print(f"PASS  {name}")
            passed += 1
        except AssertionError as e:
            print(f"FAIL  {name}: {e}")
            failed += 1
        except Exception as e:
            import traceback
            print(f"ERROR {name}: {type(e).__name__}: {e}")
            traceback.print_exc()
            failed += 1

    print(f"\n{passed}/{passed + failed} tests passés")
    return failed == 0


if __name__ == "__main__":
    success = run_all()
    sys.exit(0 if success else 1)
