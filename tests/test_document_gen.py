"""
Tests de document_gen.py — Génération des 5 conventions GENERFEU.

Vérifie :
  - 5 fichiers .docx créés (un par groupe)
  - Taille raisonnable (> 10 Ko — template non vide)
  - Noms de fichiers cohérents
  - Contenu vérifiable via python-docx (pas de tags Jinja2 résiduels)
"""

import sys
from pathlib import Path

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

from core.excel_parser import parse_formations_excel
from core.module_mapper import ModuleMapper
from core.group_builder import build_convention_groups
from core.document_gen import ConventionRenderer, ClientInfo

EXCEL_PATH = ROOT / "base de travail" / "GENERFEU - Formations et inscrits.xlsx"
OUTPUT_DIR  = ROOT / "output" / "test_conventions"

# ---------------------------------------------------------------------------
# Données GENERFEU
# ---------------------------------------------------------------------------

GENERFEU = ClientInfo(
    nom="GENERFEU",
    adresse="PARC D'ACTIVITES DES ECLAPONS\n3 CHEMIN DES ECLAPONS\n69390 VOURLES",
    representant="Gilles BERTRAND",
    fonction="Dirigeant",
    frais_mission_ht=0.0,
)

DATES = "Du 01/04/2026 au 31/10/2026"
DATE_DOC = "01/03/2026"


def _load():
    data = parse_formations_excel(EXCEL_PATH)
    mapper = ModuleMapper()
    groups = build_convention_groups(data, mapper)
    renderer = ConventionRenderer()
    return groups, renderer


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_genere_5_fichiers():
    groups, renderer = _load()
    paths = renderer.generate_all(
        groups=groups,
        client=GENERFEU,
        dates_previsionnelles=DATES,
        date_document=DATE_DOC,
        output_dir=OUTPUT_DIR,
    )
    assert len(paths) == 5, f"Attendu 5 fichiers, obtenu {len(paths)}"
    for p in paths:
        assert p.exists(), f"Fichier manquant : {p}"
        assert p.suffix == ".docx"
        print(f"  {p.name} ({p.stat().st_size // 1024} Ko)")


def test_fichiers_non_vides():
    groups, renderer = _load()
    paths = renderer.generate_all(
        groups=groups, client=GENERFEU,
        dates_previsionnelles=DATES, date_document=DATE_DOC,
        output_dir=OUTPUT_DIR,
    )
    for p in paths:
        size = p.stat().st_size
        assert size > 10_000, f"{p.name} trop petit : {size} octets"
    print(f"  Tous les fichiers > 10 Ko — OK")


def test_pas_de_tags_residuels():
    """Vérifie qu'aucun tag Jinja2 {{ }} ou {%tr %} ne subsiste dans les fichiers générés."""
    import re
    from docx import Document

    groups, renderer = _load()
    paths = renderer.generate_all(
        groups=groups, client=GENERFEU,
        dates_previsionnelles=DATES, date_document=DATE_DOC,
        output_dir=OUTPUT_DIR,
    )
    pattern = re.compile(r"\{\{|\}\}|\{%|\%\}")
    for p in paths:
        doc = Document(p)
        texts = []
        for para in doc.paragraphs:
            texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        full = " ".join(texts)
        matches = pattern.findall(full)
        assert not matches, (
            f"Tags résiduels dans {p.name} : {set(matches)}"
        )
    print("  Aucun tag Jinja2 résiduel — OK")


def test_contenu_administrateurs():
    """Vérifie le contenu de la convention Administrateurs."""
    from docx import Document

    groups, renderer = _load()
    paths = renderer.generate_all(
        groups=groups, client=GENERFEU,
        dates_previsionnelles=DATES, date_document=DATE_DOC,
        output_dir=OUTPUT_DIR,
    )
    # Premier fichier = Administrateurs (cols 5-8)
    doc = Document(paths[0])
    full_text = " ".join(
        p.text for p in doc.paragraphs
    ) + " ".join(
        cell.text for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )

    # Vérifications clés
    assert "GENERFEU" in full_text,           "Nom client manquant"
    assert "Gilles BERTRAND" in full_text,    "Représentant manquant"
    assert "Dirigeant" in full_text,          "Fonction manquante"
    assert "1-01" in full_text,               "Module 1-01 manquant"
    assert "1-03" in full_text,               "Module 1-03 manquant"
    assert "1-05" in full_text,               "Module 1-05 manquant"
    assert "1-09" in full_text,               "Module 1-09 manquant"
    assert "DUPONT" in full_text,             "Stagiaire DUPONT manquant"
    assert "NATALE" in full_text,             "Stagiaire NATALE manquant"
    assert "BERTRAND" not in _get_stagiaires_text(doc), "BERTRAND (TNS) ne doit pas être dans le tableau stagiaires"
    assert "1\xa0800" in full_text,            "Montant 1800€ manquant"
    assert "2\xa0160" in full_text,            "TTC 2160€ manquant"
    assert "01/03/2026" in full_text,         "Date document manquante"
    assert "À distance" in full_text,         "Modalité manquante"
    assert "Du 01/04/2026" in full_text,      "Dates prévisionnelles manquantes"
    print("  Administrateurs : tous les champs vérifiés — OK")


def _get_stagiaires_text(doc) -> str:
    """Extrait le texte du tableau stagiaires (Table 3)."""
    if len(doc.tables) < 4:
        return ""
    t = doc.tables[3]
    return " ".join(cell.text for row in t.rows for cell in row.cells)


def test_contenu_techniciens():
    """Vérifie la convention Techniciens : 17 stagiaires, module 1-10."""
    from docx import Document

    groups, renderer = _load()
    paths = renderer.generate_all(
        groups=groups, client=GENERFEU,
        dates_previsionnelles=DATES, date_document=DATE_DOC,
        output_dir=OUTPUT_DIR,
    )
    # Dernier fichier = Techniciens (cols 18-19)
    doc = Document(paths[4])
    full_text = " ".join(
        p.text for p in doc.paragraphs
    ) + " ".join(
        cell.text for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "1-10" in full_text,      "Module 1-10 manquant"
    assert "AGHBALOU" in full_text,  "Stagiaire AGHBALOU manquant"
    assert "TRAIY" in full_text,     "Stagiaire TRAIY manquant"

    # 17 stagiaires dans le tableau (1 en-tête + 17 lignes + 1 endfor vide)
    stagiaires_table = doc.tables[3]
    stagiaire_rows = [
        row for row in stagiaires_table.rows[1:]
        if row.cells[0].text.strip()
    ]
    assert len(stagiaire_rows) == 17, (
        f"Attendu 17 lignes stagiaires, obtenu {len(stagiaire_rows)}"
    )
    print(f"  Techniciens : {len(stagiaire_rows)} stagiaires, module 1-10 — OK")


def test_contenu_sav_gr03():
    """Vérifie SAV Gr03 : 5 stagiaires, modules 1-02 + 1-10 + 2-06, 3 DJ."""
    from docx import Document

    groups, renderer = _load()
    paths = renderer.generate_all(
        groups=groups, client=GENERFEU,
        dates_previsionnelles=DATES, date_document=DATE_DOC,
        output_dir=OUTPUT_DIR,
    )
    doc = Document(paths[3])  # SAV Gr03 = 4e groupe
    full_text = " ".join(
        cell.text for table in doc.tables
        for row in table.rows for cell in row.cells
    )
    assert "1-02" in full_text,       "Module 1-02 manquant"
    assert "1-10" in full_text,       "Module 1-10 manquant"
    assert "2-06" in full_text,       "Module 2-06 manquant"
    assert "1\xa0350" in full_text,    "Montant 1350€ manquant"
    assert "MARQUES" in full_text,    "MARQUES manquant"
    assert "VANBECELAERE" in full_text, "VANBECELAERE manquant"
    print("  SAV Gr03 : modules 1-02+1-10+2-06, 1350€ HT — OK")


def test_noms_fichiers():
    groups, renderer = _load()
    paths = renderer.generate_all(
        groups=groups, client=GENERFEU,
        dates_previsionnelles=DATES, date_document=DATE_DOC,
        output_dir=OUTPUT_DIR,
    )
    for p in paths:
        assert p.name.startswith("CONVENTION_GENERFEU_"), (
            f"Nom inattendu : {p.name}"
        )
    print("  Noms de fichiers :")
    for p in paths:
        print(f"    {p.name}")


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

def run_all():
    tests = [
        test_genere_5_fichiers,
        test_fichiers_non_vides,
        test_pas_de_tags_residuels,
        test_contenu_administrateurs,
        test_contenu_techniciens,
        test_contenu_sav_gr03,
        test_noms_fichiers,
    ]

    passed = failed = 0
    for test in tests:
        name = test.__name__
        try:
            test()
            print(f"PASS  {name}")
            passed += 1
        except AssertionError as e:
            print(f"FAIL  {name}: {e}")
            failed += 1
        except Exception as e:
            import traceback
            print(f"ERROR {name}: {type(e).__name__}: {e}")
            traceback.print_exc()
            failed += 1

    print(f"\n{passed}/{passed + failed} tests passés")
    return failed == 0


if __name__ == "__main__":
    success = run_all()
    sys.exit(0 if success else 1)
