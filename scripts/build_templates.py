"""
build_templates.py — Génère templates/convention.docx depuis le vrai document GENERFEU.

À lancer UNE FOIS (ou après modification du document source).
Injecte les tags Jinja2/docxtpl dans les bons emplacements, en conservant
toute la mise en forme (styles, polices, bordures) du document original.

Tables modifiées :
  Table 0 (1 col)  — infos client         → remplacement simple
  Table 1 (1 col)  — liste modules         → boucle {%tr for m in modules %}
  Table 2 (2 cols) — effectif / durées     → remplacement simple
  Table 3 (2 cols) — liste stagiaires      → boucle {%tr for s in stagiaires %}
  Table 4 (2 cols) — montants              → remplacement simple
  Table 5 (2 cols) — objectifs par module  → boucle {%tr for m in modules %}
"""

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent.parent
SRC  = ROOT / "base de travail" / "A signer - GENERFEU - CONVENTION Administrateurs.docx"
DEST = ROOT / "templates" / "convention.docx"


# ---------------------------------------------------------------------------
# Helpers XML
# ---------------------------------------------------------------------------

def _get_tc_list(tr_element):
    """Retourne la liste des éléments <w:tc> d'une ligne."""
    return tr_element.findall(qn("w:tc"))


def _clear_tc(tc):
    """Supprime tous les paragraphes d'une cellule sauf le premier, vide les runs."""
    paras = tc.findall(qn("w:p"))
    for p in paras[1:]:
        tc.remove(p)
    para = paras[0]
    for r in para.findall(qn("w:r")):
        para.remove(r)
    return para


def _set_tc_text(tc, text: str):
    """Vide une cellule et y place un seul run avec le texte donné."""
    para = _clear_tc(tc)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    # xml:space="preserve" pour conserver les espaces en début/fin (important pour les tags Jinja2)
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    para.append(r)


def _set_cell(cell, text: str):
    """API haut niveau : vide et remplace le texte d'une cellule."""
    _set_tc_text(cell._tc, text)


def _delete_rows_from(table, from_idx: int):
    """Supprime toutes les lignes à partir de from_idx inclus."""
    tbl = table._tbl
    rows = table.rows
    for i in range(len(rows) - 1, from_idx - 1, -1):
        tbl.remove(rows[i]._tr)


def _clone_row_with_texts(table, source_row_idx: int, texts: list[str]):
    """
    Clone la structure d'une ligne existante (avec largeurs de colonnes)
    et y injecte les textes donnés (un par cellule).
    """
    source_tr = table.rows[source_row_idx]._tr
    new_tr = deepcopy(source_tr)
    tcs = _get_tc_list(new_tr)
    for i, tc in enumerate(tcs):
        text = texts[i] if i < len(texts) else ""
        _set_tc_text(tc, text)
    table._tbl.append(new_tr)


# ---------------------------------------------------------------------------
# Construction du template
# ---------------------------------------------------------------------------

def build_convention_template():
    if not SRC.exists():
        print(f"ERREUR : document source introuvable : {SRC}", file=sys.stderr)
        sys.exit(1)

    DEST.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(SRC, DEST)
    doc = Document(DEST)

    # ------------------------------------------------------------------
    # TABLE 0 — Infos client (4 lignes × 1 colonne)
    # ------------------------------------------------------------------
    t0 = doc.tables[0]
    # Ligne 1 : nom + adresse + mention légale
    _set_cell(t0.rows[1].cells[0],
        "{{ client_nom }}\n{{ client_adresse }}\n(ci-après dénommé le bénéficiaire)")
    # Ligne 2 : représentant
    _set_cell(t0.rows[2].cells[0], "Représentée par {{ client_representant }}")
    # Ligne 3 : fonction
    _set_cell(t0.rows[3].cells[0], "Fonction : {{ client_fonction }}")

    # ------------------------------------------------------------------
    # TABLE 1 — Liste des modules (5 lignes × 1 colonne → boucle)
    #
    # Structure correcte pour docxtpl :
    #   Ligne for     : {%tr for m in modules %}   ← remplacée par {% for m in modules %}
    #   Ligne content : {{ m.code }} - ...          ← corps répété N fois
    #   Ligne endfor  : {%tr endfor %}              ← remplacée par {% endfor %}
    # ------------------------------------------------------------------
    t1 = doc.tables[1]
    # Ligne 0 : en-tête "Nom de la formation" (conserver)
    # Ligne 1 → tag de boucle seul (sera effacée, devient {% for %})
    _set_cell(t1.rows[1].cells[0], "{%tr for m in modules %}")
    # Supprimer lignes 2, 3, 4 (modules fixes) pour ne garder que ligne 1
    _delete_rows_from(t1, from_idx=2)
    # Ajouter ligne content (corps du loop — clonée de la ligne 1 pour garder le style)
    _clone_row_with_texts(t1, source_row_idx=1, texts=["{{ m.code }} - {{ m.intitule }}"])
    # Ajouter ligne endfor
    _clone_row_with_texts(t1, source_row_idx=1, texts=["{%tr endfor %}"])

    # ------------------------------------------------------------------
    # TABLE 2 — Effectif / durées / modalités (7 lignes × 2 colonnes)
    # ------------------------------------------------------------------
    t2 = doc.tables[2]
    _set_cell(t2.rows[0].cells[1], "{{ effectif }}")
    _set_cell(t2.rows[1].cells[1], "{{ nb_journees_str }}")
    _set_cell(t2.rows[2].cells[1], "{{ duree_detail }}")
    _set_cell(t2.rows[4].cells[1], "{{ modalite }}")
    # Ligne 5 : "Formation de type collectif" — laisser tel quel
    _set_cell(t2.rows[6].cells[1], "{{ dates_previsionnelles }}")

    # ------------------------------------------------------------------
    # TABLE 3 — Stagiaires (6 lignes × 2 colonnes → boucle)
    #
    # Structure correcte pour docxtpl :
    #   Ligne for     : {%tr for s in stagiaires %}  ← ligne tag seule (vide)
    #   Ligne content : {{ s.nom }} | {{ s.prenom }} ← corps répété N fois
    #   Ligne endfor  : {%tr endfor %}               ← ligne tag seule
    # ------------------------------------------------------------------
    t3 = doc.tables[3]
    # Ligne 0 : en-tête NOM | Prénom (conserver)
    # Ligne 1 → tag de boucle seul (sera effacée, devient {% for %})
    _set_cell(t3.rows[1].cells[0], "{%tr for s in stagiaires %}")
    _set_cell(t3.rows[1].cells[1], "")
    # Supprimer lignes 2-5 (stagiaires fixes de GENERFEU)
    _delete_rows_from(t3, from_idx=2)
    # Ajouter ligne content (corps du loop — clonée de la ligne 1 pour garder le style)
    _clone_row_with_texts(t3, source_row_idx=1, texts=["{{ s.nom }}", "{{ s.prenom }}"])
    # Ajouter ligne endfor
    _clone_row_with_texts(t3, source_row_idx=1, texts=["{%tr endfor %}", ""])

    # ------------------------------------------------------------------
    # TABLE 4 — Montants (5 lignes × 2 colonnes)
    # ------------------------------------------------------------------
    t4 = doc.tables[4]
    _set_cell(t4.rows[0].cells[1], "{{ montant_ht_str }}")
    _set_cell(t4.rows[1].cells[1], "{{ frais_mission_str }}")
    # Ligne 2 : vide (séparateur) — laisser
    _set_cell(t4.rows[3].cells[1], "{{ total_ht_str }}")
    _set_cell(t4.rows[4].cells[1], "{{ total_ttc_str }}")

    # ------------------------------------------------------------------
    # TABLE 5 — Objectifs pédagogiques (4 lignes × 2 colonnes → boucle)
    #
    # Structure correcte pour docxtpl :
    #   Ligne for     : {%tr for m in modules %}              ← ligne tag seule (vide)
    #   Ligne content : {{ m.code }}... | Objectifs...        ← corps répété N fois
    #   Ligne endfor  : {%tr endfor %}                        ← ligne tag seule
    # ------------------------------------------------------------------
    t5 = doc.tables[5]
    # Ligne 0 → tag de boucle seul (pas d'en-tête dans cette table)
    _set_cell(t5.rows[0].cells[0], "{%tr for m in modules %}")
    _set_cell(t5.rows[0].cells[1], "")
    # Supprimer lignes 1, 2, 3
    _delete_rows_from(t5, from_idx=1)
    # Ajouter ligne content (corps du loop — clonée de la ligne 0 pour garder le style)
    _clone_row_with_texts(t5, source_row_idx=0, texts=[
        "{{ m.code }} - {{ m.intitule }}",
        "Objectifs opérationnels : {{ m.objectif_operationnel }}\n\n"
        "A l'issue du parcours, l'apprenant sera capable de :\n"
        "{{ m.objectifs_texte }}",
    ])
    # Ajouter ligne endfor
    _clone_row_with_texts(t5, source_row_idx=0, texts=["{%tr endfor %}", ""])

    # ------------------------------------------------------------------
    # PARAGRAPHES — Date du document ("Le 26/02/2026")
    # ------------------------------------------------------------------
    import re as _re
    date_pattern = _re.compile(r"^Le\s+\d{2}/\d{2}/\d{4}$")
    for p in doc.paragraphs:
        txt = p.text.strip()
        if date_pattern.match(txt):
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = "Le {{ date_document }}"
            else:
                p.add_run("Le {{ date_document }}")
            break

    # ------------------------------------------------------------------
    # Sauvegarde
    # ------------------------------------------------------------------
    doc.save(DEST)
    print(f"Template créé : {DEST}")
    _verify_template(doc)


def _verify_template(doc):
    """Vérification rapide : liste les tags Jinja2 trouvés dans le document."""
    import re
    found = set()
    pattern = re.compile(r"\{\{[^}]+\}\}|\{%[^%]+%\}")

    for p in doc.paragraphs:
        for m in pattern.finditer(p.text):
            found.add(m.group())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for m in pattern.finditer(cell.text):
                    found.add(m.group())

    print(f"\nTags Jinja2 trouvés ({len(found)}) :")
    for tag in sorted(found):
        print(f"  {tag}")


if __name__ == "__main__":
    build_convention_template()
